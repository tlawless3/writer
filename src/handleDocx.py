import os
import docx

baseDir = os.path.abspath(os.path.join(
    os.path.dirname(__file__), '..', 'writing'))

targetDir = os.path.abspath(os.path.join(
    os.path.dirname(__file__), '..', 'outputTxtFile'))

print('basedir', baseDir)


def traverseFilesCreateDocument(folder):
    print('running')
    for root, subdirs, files in os.walk(folder):
        for curFile in files:
            if curFile.endswith('.docx'):
                curDoc = docx.Document(root + '/' + curFile)
                if os.path.isfile(targetDir + '/data.docx'):
                    targetDoc = docx.Document(targetDir + '/data.docx')
                    for paragraph in curDoc.paragraphs:
                        targetDoc.add_paragraph(paragraph.text)
                    targetDoc.save(targetDir + '/data.docx')
                else:
                    outputFile = docx.Document()
                    for paragraph in curDoc.paragraphs:
                        outputFile.add_paragraph(paragraph.text)
                    outputFile.save(targetDir + '/data.docx')


traverseFilesCreateDocument(baseDir)
